import chainlit as cl
from langchain_ollama import ChatOllama, OllamaEmbeddings
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
import hashlib
from docx import Document as DocxReader #DocxReader for no duplicated to Document

#launching models
llm = ChatOllama(
    model = 'qwen2.5-coder:7b'
)

embeddings = OllamaEmbeddings(
    model = 'nomic-embed-text'
)

def get_file_hash(filepath):
    hasher = hashlib.md5()

    with open(filepath, 'rb') as f:
        buf = f.read()
        hasher.update(buf)  

    return hasher.hexdigest()

def parse_pdf_to_docs(filename, source_name):
    reader = PdfReader(filename)
    print("reader: ", reader)
    docs = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""

        if page_text.strip():
            docs.append(
                Document(
                    page_content=f"===== PAGE {i + 1} =====\n{page_text}",
                    metadata={
                        "source": source_name,
                        "page": i + 1
                    }
                )
            )

    return docs

def parse_docx_to_docs(filename, source_name):
    reader = DocxReader(filename)
    docs = []

    paragraphs = []

    for para in reader.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    if full_text.strip():
        docs.append(
            Document(
                page_content= full_text,
                metadata={
                    "source": source_name,
                    "page": None,
                    'file_type': "docx"
                }
            )
        )

    return docs

def check_files_type(element):
    if element.mime == 'application/pdf':
        return parse_pdf_to_docs(element.path, element.name)
    
    elif element.mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return parse_docx_to_docs(element.path, element.name)
    
    else:
        return []

def chunk_docs(docs):
    splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", ". ", " ", ""],
        chunk_size=1000,
        chunk_overlap=300
    )

    return splitter.split_documents(docs)

def called_prompt(question, vectorstore, k):

    if k < 1:
        return "No documents available"

    k_min = min(8, k) # the eight-most chunks related to the question
    results = vectorstore.similarity_search(question, k = k_min)

    context = "\n\n".join([
        f"Source: {doc.metadata.get('source', 'unknown')}\n{doc.page_content}"
        for doc in results
    ])

    prompt = f"""
        You are a helpful document question-answering assistant..

        Use the context below to answer the user's question.
        The user's question may contain typos, broken English, or informal wording, or even non-English languages.
        First, check whether the input is English, then infer the user's likely intent. After this, answer based on the context.

        Context:
        {context}

        User question:
        {question}

        Instructions:
        - Answer using the context.
        - If the question is unclear, interpret it as best as possible.
        - If relevant information exists in the context, answer clearly.
        - Mention the source file and page if available.
        - If page is not available, mention only the source file.
        - Only say "I don't know" if the context has no relevant information at all.
        """
    return prompt

@cl.on_message
async def main(message: cl.Message):
    uploaded_count = 0;
    print("uploaded_count: ", uploaded_count)

    print("message.elements", message.elements)

    # Case 1: User uploads 1 PDF
    if message.elements:
        for element in message.elements:

            uploaded_hashes = cl.user_session.get("uploaded_hashes") or set()

            file_hash = get_file_hash(element.path)
            print("file_hash", file_hash)

            #check Duplicate.
            if file_hash in uploaded_hashes:
                continue

            raw_docs = check_files_type(element)

            if not raw_docs:
                await cl.Message(
                    content = f"File {element.name} is not supported now, so please convert your necessary file into PDF"
                )
            
            docs = chunk_docs(raw_docs)
            
            if docs:

                vectorstore = cl.user_session.get("vectorstore")
                all_docs = cl.user_session.get("chunks") or []

                print("vectorstore", vectorstore is None)

                if vectorstore is None:
                    print("Da vao if: ", uploaded_count)
                    vectorstore = FAISS.from_documents(
                        docs,
                        embedding= embeddings
                    )

                else:
                    print("Da vao else: ", uploaded_count)

                    
                    vectorstore.add_documents(docs)
                    print("New docs added:", len(docs))
                    print("All docs:", len(all_docs))
                    print("after added vectorstore:", vectorstore.index.ntotal)

                #extend docs into all_docs
                all_docs.extend(docs)

                # set_chunks & vectorstore for reuse purposes
                uploaded_hashes.add(file_hash)
                cl.user_session.set("uploaded_hashes", uploaded_hashes)
                cl.user_session.set("chunks", all_docs)
                cl.user_session.set("vectorstore", vectorstore)

        # in cases where users upload PDF and yield questions simultaneously
        if message.content.strip():

            question = message.content
            vectorstore = cl.user_session.get("vectorstore")
            all_docs = cl.user_session.get("chunks") or []

            prompt = called_prompt(question, vectorstore, len(all_docs))

            print("uploaded_count 2: ", uploaded_count)

            response = llm.invoke(prompt)

            await cl.Message(
                content = response.content
            ).send()

        else:
            await cl.Message(
                content = "Now ask me a question about the PDF."
            ).send()

        return
    else: 
    # Case 2: User asks a question
        vectorstore = cl.user_session.get('vectorstore')
        chunks = cl.user_session.get('chunks')

        if vectorstore is None:
            await cl.Message(
                content = '''Please upload a PDF first, as I'm the botchat that can only be useful when you upload a PDF/docs'''
            ).send()
            return
        
        print("uploaded_count 3: ", uploaded_count)
        
        prompt = called_prompt(message.content, vectorstore, len(chunks))
        response = llm.invoke(prompt)

        await cl.Message(
            content=response.content
        ).send()